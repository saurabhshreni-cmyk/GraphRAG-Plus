"""FastAPI application."""

from __future__ import annotations

import os
import uuid
from dataclasses import asdict
from pathlib import Path
from typing import Any

from fastapi import FastAPI, File, Form, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, Response

from graphrag_plus.app.config.settings import get_settings
from graphrag_plus.app.corpus.blob_store import make_blob_store
from graphrag_plus.app.corpus.seed import seed_demo_corpus
from graphrag_plus.app.evaluation.runner import evaluate_stub
from graphrag_plus.app.pipeline import GraphRAGPipeline
from graphrag_plus.app.schemas.models import (
    CorpusInfo,
    EvalResult,
    GraphResponse,
    HealthResponse,
    IngestRequest,
    IngestResponse,
    QueryRequest,
    QueryResponse,
)
from graphrag_plus.app.utils.metrics import METRICS

settings = get_settings()
# Seed the demo corpus into the active backend before the pipeline scans it.
# Idempotent: Postgres keeps it across instances; a fresh /tmp or local dir
# gets it on first boot.
seed_demo_corpus(settings.corpora_dir, make_blob_store(settings.database_url))
pipeline = GraphRAGPipeline(settings)
app = FastAPI(title="GraphRAG++")

# Edge filtering for the "important" graph view.
#   ``mentions`` is kept — it's the chunk→entity backbone, removing it leaves
#   entity nodes floating with no connections in the visualization.
#   Only drop edges that have an explicit confidence below the floor.
_LOW_VALUE_EDGES: set[str] = set()
_MIN_EDGE_CONFIDENCE = 0.6

# CORS — origins controlled by env (comma-separated). Defaults are local Vite dev servers.
_default_origins = "http://localhost:5173,http://127.0.0.1:5173"
_origins = [
    origin.strip()
    for origin in os.environ.get("GRAPHRAG_CORS_ORIGINS", _default_origins).split(",")
    if origin.strip()
]
app.add_middleware(
    CORSMiddleware,
    allow_origins=_origins or ["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.exception_handler(Exception)
async def _unhandled_exception_handler(request: Request, exc: Exception) -> JSONResponse:
    """Return a CORS-friendly JSON 500 for any unhandled error.

    Starlette's default ServerErrorMiddleware sits *outside* CORSMiddleware,
    so a bare unhandled exception yields a 500 with no ``Access-Control-
    Allow-Origin`` header — browsers then surface an opaque "Failed to fetch"
    instead of the real error. Registering this handler keeps the response
    inside the CORS layer so the frontend can read and display the message.
    """
    pipeline.logger.exception("api.unhandled_error path=%s", request.url.path)
    return JSONResponse(
        status_code=500,
        content={"detail": f"Internal error: {type(exc).__name__}"},
    )


@app.get("/health", response_model=HealthResponse)
def health() -> HealthResponse:
    """Health check."""
    return HealthResponse(
        status="ok",
        llm_enabled=settings.llm_enabled,
        graph_exists=settings.graph_path.exists(),
    )


@app.post("/ingest", response_model=IngestResponse)
def ingest(request: IngestRequest) -> IngestResponse:
    """Ingest files and URLs into a (new or existing) corpus."""
    return pipeline.ingest(
        request.file_paths,
        request.urls,
        corpus_id=request.corpus_id,
        new_corpus=request.new_corpus,
        corpus_name=request.corpus_name,
    )


@app.get("/corpora", response_model=list[CorpusInfo])
def list_corpora() -> list[CorpusInfo]:
    """List all known corpora (newest first)."""
    return [CorpusInfo(**asdict(meta)) for meta in pipeline.corpus_manager.list()]


@app.get("/corpora/active", response_model=CorpusInfo | None)
def active_corpus() -> CorpusInfo | None:
    """Return the currently active corpus, if any."""
    bundle = pipeline.corpus_manager.get_active()
    return CorpusInfo(**asdict(bundle.meta)) if bundle else None


@app.post("/corpora/{corpus_id}/select", response_model=CorpusInfo)
def select_corpus(corpus_id: str) -> CorpusInfo:
    """Switch the pipeline's active corpus to ``corpus_id``."""
    try:
        bundle = pipeline.corpus_manager.set_active(corpus_id)
    except KeyError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    return CorpusInfo(**asdict(bundle.meta))


@app.delete("/corpora/{corpus_id}")
def delete_corpus(corpus_id: str) -> dict[str, str]:
    """Permanently delete a corpus (its directory + cached state + Neo4j nodes)."""
    try:
        pipeline.corpus_manager.delete(corpus_id)
    except KeyError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    # Best-effort cleanup of the corpus's Neo4j partition; the file/blob
    # stores are already gone, so a failure here only leaves orphaned nodes.
    try:
        from graphrag_plus.app.graph.neo4j_store import get_neo4j_store

        get_neo4j_store().clear_corpus(corpus_id)
    except Exception:
        pipeline.logger.warning("neo4j.delete_cleanup_failed corpus=%s", corpus_id)
    return {"status": "deleted", "corpus_id": corpus_id}


@app.post("/query", response_model=QueryResponse)
def query(request: QueryRequest) -> QueryResponse:
    """Run question answering."""
    return pipeline.query(request)


@app.get("/graph")
def graph_snapshot(
    mode: str = "important",
    limit: int = 60,
    full_limit: int = 500,
    corpus_id: str | None = None,
    max_edges_per_node: int = 8,
    min_co_occurs: int = 2,
) -> dict[str, list[dict[str, Any]]]:
    """Return a corpus's graph for visualization.

    * ``corpus_id`` — which isolated corpus to read; defaults to the active
      corpus. Each corpus's graph is fully isolated from the others.
    * ``mode=important`` (default) — concept-centric view: top ``limit``
      entity nodes ranked by degree, with per-node edge cap and a
      co-occurrence frequency floor to keep dense graphs readable.
    * ``mode=full`` — raw graph snapshot up to ``full_limit``.
    * ``max_edges_per_node`` — only the strongest N edges per node survive
      in IMPORTANT view.
    * ``min_co_occurs`` — co-occurrence edges with weight below this are
      dropped (still meaningful relations like is_a / supports stay).
    """
    if corpus_id:
        try:
            bundle = pipeline.corpus_manager.get(corpus_id)
        except KeyError as exc:
            raise HTTPException(status_code=404, detail=str(exc)) from exc
    else:
        bundle = pipeline.corpus_manager.get_active() or pipeline._resolve_corpus(None)
    snapshot = bundle.graph_store.current_snapshot()
    nodes: list[dict[str, Any]] = snapshot.get("nodes", [])
    edges: list[dict[str, Any]] = snapshot.get("edges", [])

    if mode == "full":
        return {"nodes": nodes[:full_limit], "edges": edges[:full_limit]}

    # ---- "important" mode: filters for a readable concept graph ------------
    # Filter out: chunk-mention noise, low-confidence edges, weak co-occurrence
    # pairs (weight < min_co_occurs), and generic doc→chunk structural edges.

    def _edge_weight(edge: dict[str, Any]) -> float:
        """Strength used to rank edges. co_occurs uses ``weight``;
        meaningful relations use ``confidence``; everything else gets 0.5."""
        if edge.get("edge_type") == "co_occurs":
            return float(edge.get("weight", 1))
        conf = edge.get("confidence")
        if isinstance(conf, (int | float)):
            return float(conf)
        return 0.5

    def _keep_edge(edge: dict[str, Any]) -> bool:
        et = edge.get("edge_type")
        # Drop the chunk → entity boilerplate and doc → chunk backbone — this
        # view is concept-centric.
        if et in {"mentions", "contains"}:
            return False
        # Drop weak co-occurrence pairs that are below the frequency floor.
        if et == "co_occurs" and float(edge.get("weight", 0)) < min_co_occurs:
            return False
        conf = edge.get("confidence")
        return not (isinstance(conf, (int | float)) and conf < _MIN_EDGE_CONFIDENCE)

    filtered_edges = [e for e in edges if _keep_edge(e)]

    degree: dict[str, int] = {}
    for edge in filtered_edges:
        degree[edge["source"]] = degree.get(edge["source"], 0) + 1
        degree[edge["target"]] = degree.get(edge["target"], 0) + 1

    # Concept-centric view: drop doc + chunk nodes entirely. The graph
    # answers "what concepts in this corpus relate to each other?".
    entity_nodes = [n for n in nodes if n.get("node_type") not in {"Document", "Chunk"}]
    entity_nodes.sort(key=lambda n: degree.get(n["id"], 0), reverse=True)

    keep_entities = entity_nodes[: max(1, limit)]
    keep_ids = {n["id"] for n in keep_entities}

    pruned_nodes: list[dict[str, Any]] = [dict(n) for n in nodes if n["id"] in keep_ids]
    candidate_edges = [e for e in filtered_edges if e["source"] in keep_ids and e["target"] in keep_ids]

    # Per-node edge cap: keep at most ``max_edges_per_node`` edges
    # incident to each node, picking the strongest ones first. Two-sided
    # cap — both endpoints must still have budget.
    candidate_edges.sort(key=_edge_weight, reverse=True)
    edge_count: dict[str, int] = {}
    capped_edges: list[dict[str, Any]] = []
    for edge in candidate_edges:
        s, t = edge["source"], edge["target"]
        if edge_count.get(s, 0) >= max_edges_per_node or edge_count.get(t, 0) >= max_edges_per_node:
            continue
        capped_edges.append(edge)
        edge_count[s] = edge_count.get(s, 0) + 1
        edge_count[t] = edge_count.get(t, 0) + 1

    # Annotate each node with its degree (after capping) so the frontend
    # can size by importance.
    final_degree: dict[str, int] = {}
    for edge in capped_edges:
        final_degree[edge["source"]] = final_degree.get(edge["source"], 0) + 1
        final_degree[edge["target"]] = final_degree.get(edge["target"], 0) + 1
    for n in pruned_nodes:
        n["degree"] = final_degree.get(n["id"], 0)

    return {"nodes": pruned_nodes, "edges": capped_edges}


@app.get("/graph/{node_id}", response_model=GraphResponse)
def graph(node_id: str) -> GraphResponse:
    """Return neighborhood for node."""
    neighbors = pipeline.graph_store.neighbors(node_id)
    if not neighbors:
        raise HTTPException(status_code=404, detail=f"Node '{node_id}' not found or has no neighbors.")
    return GraphResponse(node_id=node_id, neighbors=neighbors)


@app.get("/evaluate", response_model=EvalResult)
def evaluate() -> EvalResult:
    """Run benchmark evaluation."""
    result = evaluate_stub(settings.reports_dir, settings.data_dir / "benchmark.json")
    return EvalResult(metrics=result["metrics"], report_path=result["report_path"])


@app.get("/metrics")
def metrics() -> Response:
    """Prometheus-compatible metrics."""
    body, content_type = METRICS.render()
    return Response(content=body, media_type=content_type)


# =========================================================================
# Neo4j graph exploration + upload endpoints (frontend v2)
# =========================================================================


def _neo4j():
    from graphrag_plus.app.graph.neo4j_store import get_neo4j_store

    return get_neo4j_store()


@app.get("/graph/{corpus_id}/full")
def graph_full(corpus_id: str, max_nodes: int = 500) -> dict[str, Any]:
    """Full Neo4j graph for a corpus: Entity + Chunk nodes, all edges.

    Capped at ``max_nodes`` (highest connection-count entities win). Chunk
    nodes are included only when they fit inside the remaining budget.
    """
    store = _neo4j()
    entity_rows = store._run(
        """
        MATCH (e:Entity {corpus_id: $cid})
        OPTIONAL MATCH (e)-[r]-()
        WITH e, count(r) AS degree
        RETURN e.name AS id, e.name AS label, coalesce(e.type, 'OTHER') AS type,
               coalesce(e.description, '') AS description, degree AS connection_count
        ORDER BY degree DESC
        LIMIT $max_nodes
        """,
        cid=corpus_id,
        max_nodes=max(1, min(max_nodes, 500)),
    )
    if not entity_rows:
        raise HTTPException(status_code=404, detail=f"No graph data for corpus '{corpus_id}'")
    keep = {row["id"] for row in entity_rows}

    edge_rows = store._run(
        """
        MATCH (a:Entity {corpus_id: $cid})-[r:RELATES_TO]->(b:Entity {corpus_id: $cid})
        RETURN a.name AS source, b.name AS target, r.relation AS relation,
               coalesce(r.confidence, 0.5) AS weight
        """,
        cid=corpus_id,
    )
    edges = [e for e in edge_rows if e["source"] in keep and e["target"] in keep]

    type_counts: dict[str, int] = {}
    for row in entity_rows:
        type_counts[row["type"]] = type_counts.get(row["type"], 0) + 1
    total_entities = store._run(
        "MATCH (e:Entity {corpus_id: $cid}) RETURN count(e) AS c", cid=corpus_id
    )
    total_edges = store._run(
        "MATCH (:Entity {corpus_id: $cid})-[r:RELATES_TO]->(:Entity {corpus_id: $cid}) RETURN count(r) AS c",
        cid=corpus_id,
    )
    return {
        "nodes": entity_rows,
        "edges": edges,
        "stats": {
            "total_nodes": int(total_entities[0]["c"]) if total_entities else len(entity_rows),
            "total_edges": int(total_edges[0]["c"]) if total_edges else len(edges),
            "entity_types": type_counts,
            "returned_nodes": len(entity_rows),
        },
    }


@app.get("/graph/{corpus_id}/query-path")
def graph_query_path(corpus_id: str, q: str) -> dict[str, Any]:
    """Nodes + edges a query would traverse (spaCy entities only — fast).

    Used by the frontend to highlight the retrieval path over the graph.
    """
    from graphrag_plus.app.extraction.extractor import extract

    result = extract(q, use_llm=False)
    seed_names = [e.name for e in result.entities]
    if not seed_names:
        # Fall back to content tokens for short keyword queries.
        import re as _re

        seed_names = [t for t in _re.findall(r"[A-Za-z0-9]{3,}", q)][:5]

    store = _neo4j()
    matched = store._run(
        """
        UNWIND $names AS name
        MATCH (e:Entity {corpus_id: $cid})
        WHERE toLower(e.name) = toLower(name) OR toLower(e.name) CONTAINS toLower(name)
        RETURN DISTINCT e.name AS id, coalesce(e.type, 'OTHER') AS type
        LIMIT 25
        """,
        names=seed_names,
        cid=corpus_id,
    )
    seed_ids = [row["id"] for row in matched]
    edges = (
        store._run(
            """
            UNWIND $names AS name
            MATCH (a:Entity {name: name, corpus_id: $cid})-[r:RELATES_TO]-(b:Entity {corpus_id: $cid})
            RETURN DISTINCT a.name AS source, b.name AS target, r.relation AS relation
            LIMIT 100
            """,
            names=seed_ids,
            cid=corpus_id,
        )
        if seed_ids
        else []
    )
    neighbour_ids = sorted({e["target"] for e in edges} | {e["source"] for e in edges} - set(seed_ids))
    return {
        "query": q,
        "extracted_entities": seed_names,
        "seed_nodes": seed_ids,
        "neighbor_nodes": neighbour_ids,
        "edges": edges,
    }


@app.get("/graph/{corpus_id}/entity/{entity_name}")
def graph_entity(corpus_id: str, entity_name: str) -> dict[str, Any]:
    """One entity's metadata, neighbours (with relation types), and chunks."""
    store = _neo4j()
    meta = store._run(
        """
        MATCH (e:Entity {corpus_id: $cid})
        WHERE toLower(e.name) = toLower($name)
        RETURN e.name AS name, coalesce(e.type, 'OTHER') AS type,
               coalesce(e.description, '') AS description,
               coalesce(e.confidence, 0.0) AS confidence
        LIMIT 1
        """,
        cid=corpus_id,
        name=entity_name,
    )
    if not meta:
        raise HTTPException(status_code=404, detail=f"Entity '{entity_name}' not found in corpus")
    canonical = meta[0]["name"]
    neighbours = store._run(
        """
        MATCH (e:Entity {name: $name, corpus_id: $cid})-[r:RELATES_TO]-(nbr:Entity)
        RETURN DISTINCT nbr.name AS name, coalesce(nbr.type, 'OTHER') AS type,
               r.relation AS relation,
               CASE WHEN startNode(r) = e THEN 'out' ELSE 'in' END AS direction
        LIMIT 50
        """,
        name=canonical,
        cid=corpus_id,
    )
    chunks = store._run(
        """
        MATCH (e:Entity {name: $name, corpus_id: $cid})-[:MENTIONED_IN]->(c:Chunk)
        RETURN c.chunk_id AS chunk_id, c.text AS text, c.source AS source
        LIMIT 20
        """,
        name=canonical,
        cid=corpus_id,
    )
    return {"entity": meta[0], "neighbors": neighbours, "chunks": chunks}


@app.get("/graph/{corpus_id}/stats")
def graph_stats(corpus_id: str) -> dict[str, Any]:
    """Corpus graph statistics: type histograms + most connected entities."""
    store = _neo4j()
    nodes_by_type = store._run(
        """
        MATCH (e:Entity {corpus_id: $cid})
        RETURN coalesce(e.type, 'OTHER') AS type, count(e) AS count
        ORDER BY count DESC
        """,
        cid=corpus_id,
    )
    rels_by_type = store._run(
        """
        MATCH (:Entity {corpus_id: $cid})-[r:RELATES_TO]->(:Entity {corpus_id: $cid})
        RETURN r.relation AS relation, count(r) AS count
        ORDER BY count DESC
        LIMIT 25
        """,
        cid=corpus_id,
    )
    chunk_count = store._run(
        "MATCH (c:Chunk {corpus_id: $cid}) RETURN count(c) AS c", cid=corpus_id
    )
    top_entities = store._run(
        """
        MATCH (e:Entity {corpus_id: $cid})
        OPTIONAL MATCH (e)-[r]-()
        WITH e, count(r) AS degree
        RETURN e.name AS name, coalesce(e.type, 'OTHER') AS type, degree
        ORDER BY degree DESC
        LIMIT 10
        """,
        cid=corpus_id,
    )
    if not nodes_by_type and not chunk_count:
        raise HTTPException(status_code=404, detail=f"No graph data for corpus '{corpus_id}'")
    return {
        "corpus_id": corpus_id,
        "nodes_by_type": {row["type"]: int(row["count"]) for row in nodes_by_type},
        "relationships_by_type": {str(row["relation"]): int(row["count"]) for row in rels_by_type},
        "total_chunks": int(chunk_count[0]["c"]) if chunk_count else 0,
        "top_connected_entities": top_entities,
    }


_UPLOAD_EXTENSIONS = {".pdf", ".txt", ".md", ".docx", ".html", ".htm"}


def _extract_docx_text(path: Path) -> str:
    """Plain-text extraction from a DOCX (paragraphs + table cells)."""
    import docx

    document = docx.Document(str(path))
    parts = [para.text for para in document.paragraphs if para.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


@app.post("/ingest/file")
def ingest_file(
    file: UploadFile = File(...),
    corpus_id: str | None = Form(default=None),
    corpus_name: str | None = Form(default=None),
) -> dict[str, Any]:
    """Upload one document (PDF / TXT / MD / DOCX / HTML) and ingest it.

    Creates a new corpus unless an existing ``corpus_id`` is supplied.
    Ingestion runs synchronously; the response carries the final counts.
    """
    suffix = Path(file.filename or "upload.txt").suffix.lower()
    if suffix not in _UPLOAD_EXTENSIONS:
        raise HTTPException(
            status_code=415,
            detail=f"Unsupported file type '{suffix}'. Allowed: {sorted(_UPLOAD_EXTENSIONS)}",
        )
    job_id = f"job_{uuid.uuid4().hex[:12]}"
    upload_dir = settings.temp_dir / "uploads"
    upload_dir.mkdir(parents=True, exist_ok=True)
    safe_stem = "".join(ch for ch in Path(file.filename or "upload").stem if ch.isalnum() or ch in "-_ ")[:60]
    target = upload_dir / f"{job_id}_{safe_stem}{suffix}"
    try:
        target.write_bytes(file.file.read())
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to store upload: {exc}") from exc

    # DOCX has no loader support — convert to plain text alongside.
    ingest_path = target
    if suffix == ".docx":
        try:
            text = _extract_docx_text(target)
        except Exception as exc:
            raise HTTPException(status_code=422, detail=f"Could not read DOCX: {exc}") from exc
        if not text.strip():
            raise HTTPException(status_code=422, detail="DOCX contained no extractable text")
        ingest_path = target.with_suffix(".txt")
        ingest_path.write_text(text, encoding="utf-8")

    response = pipeline.ingest(
        [str(ingest_path)],
        [],
        corpus_id=corpus_id,
        new_corpus=corpus_id is None,
        corpus_name=corpus_name or safe_stem or None,
    )
    return {
        "job_id": job_id,
        "corpus_id": response.corpus_id,
        "status": "completed",
        "filename": file.filename,
        "documents": response.documents,
        "chunks": response.chunks,
        "entities": response.entities,
        "relations": response.relations,
        "warnings": response.warnings,
    }


@app.get("/corpora/{corpus_id}/chunks")
def corpus_chunks(corpus_id: str, page: int = 1, page_size: int = 20) -> dict[str, Any]:
    """Paginated chunk listing for a corpus (ingestion debugging)."""
    try:
        bundle = pipeline.corpus_manager.get(corpus_id)
    except KeyError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    chunks = bundle.retrieval.chunks
    page = max(1, page)
    page_size = max(1, min(page_size, 100))
    start = (page - 1) * page_size
    rows = [
        {
            "chunk_id": chunk.chunk_id,
            "doc_id": chunk.doc_id,
            "text": chunk.text,
            "start": chunk.start,
            "end": chunk.end,
        }
        for chunk in chunks[start : start + page_size]
    ]
    return {
        "corpus_id": corpus_id,
        "page": page,
        "page_size": page_size,
        "total_chunks": len(chunks),
        "chunks": rows,
    }
